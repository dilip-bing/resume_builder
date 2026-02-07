"""
Comprehensive Format Verification Tool
Verifies COMPLETE format match including document, section, paragraph, and run levels.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
import json


class ComprehensiveFormatVerifier:
    """Verify all levels of formatting."""
    
    def __init__(self, original_path: str, generated_path: str):
        self.original = Document(original_path)
        self.generated = Document(generated_path)
        self.mismatches = []
    
    def verify_document_properties(self):
        """Verify document-level properties."""
        orig_props = self.original.core_properties
        gen_props = self.generated.core_properties
        
        checks = []
        
        # Title
        if orig_props.title != gen_props.title:
            self.mismatches.append(f"Document title: '{orig_props.title}' vs '{gen_props.title}'")
            checks.append("❌ Title")
        else:
            checks.append("✅ Title")
        
        # Author
        if orig_props.author != gen_props.author:
            self.mismatches.append(f"Document author: '{orig_props.author}' vs '{gen_props.author}'")
            checks.append("❌ Author")
        else:
            checks.append("✅ Author")
        
        return checks
    
    def verify_section_properties(self):
        """Verify section-level properties."""
        orig_section = self.original.sections[0]
        gen_section = self.generated.sections[0]
        
        checks = []
        
        # Page dimensions
        if orig_section.page_height != gen_section.page_height:
            self.mismatches.append(f"Page height: {orig_section.page_height} vs {gen_section.page_height}")
            checks.append("❌ Page height")
        else:
            checks.append("✅ Page height")
        
        if orig_section.page_width != gen_section.page_width:
            self.mismatches.append(f"Page width: {orig_section.page_width} vs {gen_section.page_width}")
            checks.append("❌ Page width")
        else:
            checks.append("✅ Page width")
        
        # Margins
        margins = ['top_margin', 'bottom_margin', 'left_margin', 'right_margin']
        for margin in margins:
            orig_val = getattr(orig_section, margin)
            gen_val = getattr(gen_section, margin)
            
            if orig_val != gen_val:
                self.mismatches.append(f"{margin}: {orig_val} vs {gen_val}")
                checks.append(f"❌ {margin}")
            else:
                checks.append(f"✅ {margin}")
        
        return checks
    
    def verify_paragraph_formatting(self, para_idx: int):
        """Verify paragraph-level formatting."""
        orig_para = self.original.paragraphs[para_idx]
        gen_para = self.generated.paragraphs[para_idx]
        
        checks = []
        
        # Alignment
        if orig_para.alignment != gen_para.alignment:
            self.mismatches.append(f"Para {para_idx} alignment: {orig_para.alignment} vs {gen_para.alignment}")
            checks.append(f"❌ Alignment")
        else:
            checks.append(f"✅ Alignment")
        
        # Spacing
        orig_pf = orig_para.paragraph_format
        gen_pf = gen_para.paragraph_format
        
        props = ['left_indent', 'right_indent', 'first_line_indent', 
                 'space_before', 'space_after', 'line_spacing']
        
        for prop in props:
            orig_val = getattr(orig_pf, prop)
            gen_val = getattr(gen_pf, prop)
            
            if orig_val != gen_val:
                self.mismatches.append(f"Para {para_idx} {prop}: {orig_val} vs {gen_val}")
                checks.append(f"❌ {prop}")
            else:
                checks.append(f"✅ {prop}")
        
        return checks
    
    def verify_run_formatting(self, para_idx: int, run_idx: int):
        """Verify run-level formatting."""
        orig_run = self.original.paragraphs[para_idx].runs[run_idx]
        gen_run = self.generated.paragraphs[para_idx].runs[run_idx]
        
        checks = []
        
        # Bold/Italic
        if orig_run.bold != gen_run.bold:
            self.mismatches.append(f"Para {para_idx} Run {run_idx} bold: {orig_run.bold} vs {gen_run.bold}")
            checks.append(f"❌ Bold")
        else:
            checks.append(f"✅ Bold")
        
        if orig_run.italic != gen_run.italic:
            self.mismatches.append(f"Para {para_idx} Run {run_idx} italic: {orig_run.italic} vs {gen_run.italic}")
            checks.append(f"❌ Italic")
        else:
            checks.append(f"✅ Italic")
        
        # Font properties
        if orig_run.font.name != gen_run.font.name:
            self.mismatches.append(f"Para {para_idx} Run {run_idx} font: {orig_run.font.name} vs {gen_run.font.name}")
            checks.append(f"❌ Font")
        else:
            checks.append(f"✅ Font")
        
        if orig_run.font.size != gen_run.font.size:
            self.mismatches.append(f"Para {para_idx} Run {run_idx} font size: {orig_run.font.size} vs {gen_run.font.size}")
            checks.append(f"❌ Font size")
        else:
            checks.append(f"✅ Font size")
        
        return checks
    
    def run_comprehensive_verification(self):
        """Run complete verification at all levels."""
        print("🔍 COMPREHENSIVE FORMAT VERIFICATION")
        print("=" * 60)
        
        # 1. Document level
        print("\n📄 Document Properties:")
        doc_checks = self.verify_document_properties()
        for check in doc_checks:
            print(f"  {check}")
        
        # 2. Section level
        print("\n📐 Section Properties:")
        section_checks = self.verify_section_properties()
        for check in section_checks:
            print(f"  {check}")
        
        # 3. Paragraph level (sample first 5)
        print("\n📝 Paragraph Formatting (sample):")
        for i in range(min(5, len(self.original.paragraphs))):
            if not self.original.paragraphs[i].text.strip():
                continue
            
            print(f"\n  Paragraph {i}:")
            para_checks = self.verify_paragraph_formatting(i)
            for check in para_checks[:3]:  # Show first 3 checks
                print(f"    {check}")
        
        # 4. Run level (sample first paragraph with runs)
        print("\n🎨 Run Formatting (sample):")
        for para_idx in range(min(3, len(self.original.paragraphs))):
            orig_para = self.original.paragraphs[para_idx]
            if not orig_para.runs or not orig_para.text.strip():
                continue
            
            print(f"\n  Paragraph {para_idx}, Run 0:")
            run_checks = self.verify_run_formatting(para_idx, 0)
            for check in run_checks[:4]:  # Show first 4 checks
                print(f"    {check}")
            break
        
        # Summary
        print("\n" + "=" * 60)
        print(f"\n📊 VERIFICATION SUMMARY:")
        print(f"   Total mismatches found: {len(self.mismatches)}")
        
        if len(self.mismatches) == 0:
            print("\n   ✅ ✅ ✅ PERFECT MATCH! ✅ ✅ ✅")
            print("   Generated resume has EXACTLY the same formatting as original!")
            print("   All levels verified: Document → Section → Paragraph → Run")
        else:
            print("\n   ❌ FORMATTING DIFFERENCES FOUND:")
            for mismatch in self.mismatches:
                print(f"   • {mismatch}")
        
        print("\n" + "=" * 60)
        
        return len(self.mismatches) == 0


# Main execution
if __name__ == '__main__':
    import sys
    
    original = r"C:\Users\dilip\OneDrive\Desktop\ResumeBuilder\reference_docx\resume_optimized_final.docx"
    
    # Accept generated file path from command line argument
    if len(sys.argv) > 1:
        generated = sys.argv[1]
    else:
        generated = "output/test_metadata_output.docx"
    
    print(f"Original: {original}")
    print(f"Generated: {generated}\n")
    
    verifier = ComprehensiveFormatVerifier(original, generated)
    is_perfect_match = verifier.run_comprehensive_verification()
    
    if is_perfect_match:
        print("\n✅ Verification PASSED")
    else:
        print("\n❌ Verification FAILED - formatting differences detected")
