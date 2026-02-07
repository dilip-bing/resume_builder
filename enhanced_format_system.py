"""
Enhanced Format Preservation System with Metadata Storage
Saves ALL formatting properties to a metadata JSON file for perfect replication.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import shutil
from typing import Dict, Any
from pathlib import Path


class FormatMetadata:
    """Extract and store complete formatting metadata."""
    
    def __init__(self, docx_path: str):
        self.doc = Document(docx_path)
        self.docx_path = docx_path
    
    def extract_complete_format_metadata(self) -> Dict[str, Any]:
        """
        Extract ALL formatting metadata from document.
        This creates a complete format map that can be used to replicate formatting.
        """
        metadata = {
            "source_document": self.docx_path,
            "document_properties": self.extract_document_properties(),
            "section_properties": self.extract_section_properties(),
            "paragraph_formats": {},
            "run_formats": {}
        }
        
        # Extract paragraph and run formatting for each paragraph
        for para_idx, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():
                continue
            
            # Paragraph format
            metadata["paragraph_formats"][para_idx] = self.extract_paragraph_format(para)
            
            # Run formats
            metadata["run_formats"][para_idx] = []
            for run_idx, run in enumerate(para.runs):
                metadata["run_formats"][para_idx].append(self.extract_run_format(run))
        
        return metadata
    
    def extract_document_properties(self) -> Dict[str, Any]:
        """Extract document-level properties."""
        core_props = self.doc.core_properties
        
        return {
            "title": core_props.title,
            "author": core_props.author,
            "subject": core_props.subject,
            "keywords": core_props.keywords
        }
    
    def extract_section_properties(self) -> Dict[str, Any]:
        """Extract section properties (margins, page size)."""
        section = self.doc.sections[0]
        
        return {
            "page_height": section.page_height,
            "page_width": section.page_width,
            "top_margin": section.top_margin,
            "bottom_margin": section.bottom_margin,
            "left_margin": section.left_margin,
            "right_margin": section.right_margin,
            "header_distance": section.header_distance,
            "footer_distance": section.footer_distance,
            "orientation": str(section.orientation)
        }
    
    def extract_paragraph_format(self, para) -> Dict[str, Any]:
        """Extract complete paragraph formatting."""
        pf = para.paragraph_format
        
        return {
            "alignment": str(para.alignment) if para.alignment else None,
            "left_indent": pf.left_indent,
            "right_indent": pf.right_indent,
            "first_line_indent": pf.first_line_indent,
            "space_before": pf.space_before,
            "space_after": pf.space_after,
            "line_spacing": pf.line_spacing,
            "line_spacing_rule": str(pf.line_spacing_rule) if pf.line_spacing_rule else None,
            "keep_together": pf.keep_together,
            "keep_with_next": pf.keep_with_next,
            "page_break_before": pf.page_break_before,
            "widow_control": pf.widow_control
        }
    
    def extract_run_format(self, run) -> Dict[str, Any]:
        """Extract complete run formatting."""
        return {
            "text": run.text,
            "bold": run.bold,
            "italic": run.italic,
            "underline": str(run.underline) if run.underline else None,
            "font_name": run.font.name,
            "font_size": run.font.size,
            "font_color_rgb": list(run.font.color.rgb) if run.font.color.rgb else None,
            "all_caps": run.font.all_caps,
            "small_caps": run.font.small_caps,
            "strike": run.font.strike,
            "double_strike": run.font.double_strike,
            "subscript": run.font.subscript,
            "superscript": run.font.superscript
        }
    
    def save_metadata(self, output_path: str = "metadata/format_metadata.json"):
        """Save complete format metadata to JSON file."""
        metadata = self.extract_complete_format_metadata()
        
        # Ensure directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Custom JSON encoder for non-serializable types
        def convert_value(obj):
            if obj is None:
                return None
            if isinstance(obj, (int, float, str, bool)):
                return obj
            return str(obj)
        
        # Convert metadata
        cleaned_metadata = json.loads(
            json.dumps(metadata, default=convert_value)
        )
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(cleaned_metadata, f, indent=2, ensure_ascii=False)
        
        print(f"✅ Saved format metadata: {output_path}")
        self.print_metadata_summary(cleaned_metadata)
        
        return output_path
    
    def print_metadata_summary(self, metadata: Dict[str, Any]):
        """Print summary of extracted metadata."""
        print(f"\n📊 Format Metadata Summary:")
        print(f"   Document: {metadata['source_document']}")
        print(f"   Paragraphs with formatting: {len(metadata['paragraph_formats'])}")
        print(f"   Total run formats: {sum(len(runs) for runs in metadata['run_formats'].values())}")
        
        section = metadata['section_properties']
        print(f"\n📄 Page Setup:")
        print(f"   Page size: {section['page_width']} x {section['page_height']}")
        print(f"   Margins: T={section['top_margin']}, B={section['bottom_margin']}, "
              f"L={section['left_margin']}, R={section['right_margin']}")


class EnhancedFormatBuilder:
    """
    Build resume using stored format metadata for perfect replication.
    """
    
    def __init__(self, original_resume_path: str, metadata_path: str = "metadata/format_metadata.json"):
        self.original_resume_path = original_resume_path
        self.original_doc = Document(original_resume_path)
        
        # Load format metadata
        with open(metadata_path, 'r', encoding='utf-8') as f:
            self.format_metadata = json.load(f)
        
        print(f"✅ Loaded format metadata from: {metadata_path}")
    
    def apply_section_properties(self, target_doc: Document):
        """Apply section properties from metadata."""
        section_meta = self.format_metadata['section_properties']
        target_section = target_doc.sections[0]
        
        if section_meta['page_height']:
            target_section.page_height = section_meta['page_height']
        if section_meta['page_width']:
            target_section.page_width = section_meta['page_width']
        if section_meta['top_margin']:
            target_section.top_margin = section_meta['top_margin']
        if section_meta['bottom_margin']:
            target_section.bottom_margin = section_meta['bottom_margin']
        if section_meta['left_margin']:
            target_section.left_margin = section_meta['left_margin']
        if section_meta['right_margin']:
            target_section.right_margin = section_meta['right_margin']
    
    def apply_paragraph_format_from_metadata(self, para, para_idx: int):
        """Apply paragraph formatting from metadata."""
        if str(para_idx) not in self.format_metadata['paragraph_formats']:
            return
        
        meta = self.format_metadata['paragraph_formats'][str(para_idx)]
        pf = para.paragraph_format
        
        # Apply all properties
        if meta['left_indent']:
            pf.left_indent = meta['left_indent']
        if meta['right_indent']:
            pf.right_indent = meta['right_indent']
        if meta['first_line_indent']:
            pf.first_line_indent = meta['first_line_indent']
        if meta['space_before']:
            pf.space_before = meta['space_before']
        if meta['space_after']:
            pf.space_after = meta['space_after']
        if meta['line_spacing']:
            pf.line_spacing = meta['line_spacing']
        
        pf.keep_together = meta.get('keep_together', False)
        pf.keep_with_next = meta.get('keep_with_next', False)
        pf.page_break_before = meta.get('page_break_before', False)
        pf.widow_control = meta.get('widow_control', True)
    
    def apply_run_format_from_metadata(self, run, para_idx: int, run_idx: int):
        """Apply run formatting from metadata."""
        if str(para_idx) not in self.format_metadata['run_formats']:
            return
        
        run_formats = self.format_metadata['run_formats'][str(para_idx)]
        if run_idx >= len(run_formats):
            return
        
        meta = run_formats[run_idx]
        
        # Apply all properties
        run.bold = meta['bold']
        run.italic = meta['italic']
        
        if meta['font_name']:
            run.font.name = meta['font_name']
        if meta['font_size']:
            run.font.size = meta['font_size']
        if meta['font_color_rgb']:
            run.font.color.rgb = RGBColor(*meta['font_color_rgb'])
        
        run.font.all_caps = meta.get('all_caps', False)
        run.font.small_caps = meta.get('small_caps', False)
        run.font.strike = meta.get('strike', False)
        run.font.double_strike = meta.get('double_strike', False)
        run.font.subscript = meta.get('subscript', False)
        run.font.superscript = meta.get('superscript', False)
    
    def replace_paragraph_with_metadata(self, para_idx: int, new_text: str, target_doc: Document):
        """Replace paragraph content using stored metadata."""
        target_para = target_doc.paragraphs[para_idx]
        
        # Clear existing runs
        for run in target_para.runs:
            run.text = ""
        
        # Apply paragraph format from metadata
        self.apply_paragraph_format_from_metadata(target_para, para_idx)
        
        # Create run with text
        if target_para.runs:
            target_run = target_para.runs[0]
        else:
            target_run = target_para.add_run()
        
        target_run.text = new_text
        
        # Apply run format from metadata (use first run's format)
        self.apply_run_format_from_metadata(target_run, para_idx, 0)
    
    def replace_paragraph_multirun_with_metadata(self, para_idx: int, text_parts: list, target_doc: Document):
        """Replace paragraph with multiple runs preserving original formatting for each part.
        
        Args:
            para_idx: Paragraph index
            text_parts: List of tuples [(text, run_indices_to_use), ...]
                       e.g., [("University Name", [0]), ("spacing", [1]), ("Degree", [2]), ...]
        """
        target_para = target_doc.paragraphs[para_idx]
        
        # Clear existing runs
        for run in target_para.runs:
            run.text = ""
        
        # Apply paragraph format from metadata
        self.apply_paragraph_format_from_metadata(target_para, para_idx)
        
        # Get original run metadata
        original_runs = self.format_metadata['run_formats'].get(str(para_idx), [])
        
        # Recreate runs with proper formatting
        run_counter = 0
        for text, run_idx in text_parts:
            if not text:
                continue
            
            # Reuse existing cleared run or create new one
            if run_counter < len(target_para.runs):
                new_run = target_para.runs[run_counter]
            else:
                new_run = target_para.add_run()
            
            new_run.text = text
            run_counter += 1
            
            # Apply formatting from the specified original run
            if run_idx < len(original_runs):
                meta = original_runs[run_idx]
                new_run.bold = meta['bold']
                new_run.italic = meta['italic']
                
                if meta['font_name']:
                    new_run.font.name = meta['font_name']
                if meta['font_size']:
                    new_run.font.size = meta['font_size']
                if meta['font_color_rgb']:
                    new_run.font.color.rgb = RGBColor(*meta['font_color_rgb'])
                
                new_run.font.all_caps = meta.get('all_caps', False)
                new_run.font.small_caps = meta.get('small_caps', False)
                new_run.font.strike = meta.get('strike', False)
                new_run.font.double_strike = meta.get('double_strike', False)
                new_run.font.subscript = meta.get('subscript', False)
                new_run.font.superscript = meta.get('superscript', False)


    
    def replace_skill_line_with_metadata(self, para_idx: int, label: str, value: str, target_doc: Document):
        """Replace skill line using stored metadata for exact formatting."""
        target_para = target_doc.paragraphs[para_idx]
        
        # Clear runs
        for run in target_para.runs:
            run.text = ""
        
        # Apply paragraph format
        self.apply_paragraph_format_from_metadata(target_para, para_idx)
        
        # Run 0: Label
        if target_para.runs:
            run0 = target_para.runs[0]
        else:
            run0 = target_para.add_run()
        
        run0.text = f"{label}: "
        self.apply_run_format_from_metadata(run0, para_idx, 0)
        
        # Run 1: Value
        if len(target_para.runs) > 1:
            run1 = target_para.runs[1]
        else:
            run1 = target_para.add_run()
        
        run1.text = value
        self.apply_run_format_from_metadata(run1, para_idx, 1)
    
    def build_resume_from_json(self, json_data: Dict[str, Any], output_path: str = "output/resume.docx") -> str:
        """
        Build resume using format metadata for perfect replication.
        """
        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Copy original
        shutil.copy2(self.original_resume_path, output_path)
        print(f"✅ Copied original to: {output_path}")
        
        # Open and apply metadata-driven formatting
        target_doc = Document(output_path)
        
        # Apply section properties
        self.apply_section_properties(target_doc)
        
        changes = []
        
        # Apply all edits using metadata
        # Personal
        if "name" in json_data.get("personal", {}):
            name_data = json_data["personal"]["name"]
            para_idx = name_data["paragraph_index"]
            self.replace_paragraph_with_metadata(para_idx, name_data["value"], target_doc)
            changes.append(f"NAME (para {para_idx})")
        
        if "contact_line" in json_data.get("personal", {}):
            contact_data = json_data["personal"]["contact_line"]
            para_idx = contact_data["paragraph_index"]
            parts = contact_data["parts"]
            contact_line = f"{parts['location']} | {parts['phone']} | {parts['email']} | {parts['linkedin']} | {parts['portfolio']}"
            self.replace_paragraph_with_metadata(para_idx, contact_line, target_doc)
            changes.append(f"CONTACT (para {para_idx})")
        
        # Education
        for idx, edu in enumerate(json_data.get("education", [])):
            if "institution_line" in edu:
                inst_data = edu["institution_line"]
                para_idx = inst_data["paragraph_index"]
                university = edu.get("university", "")
                degree = edu.get("degree", "")
                dates = edu.get("dates", "")
                
                # Education 1 (para 4): Full line with university + degree + dates
                # Education 2 (para 7): Only university name (skip this, use para 8 instead)
                if idx == 0:
                    # First education: Write full line with multirun formatting
                    text_parts = [
                        (university, 0),  # Bold
                        ("                                                          ", 1),  # Spacing
                        (degree, 2),  # Italic
                        ("\t                                                      \t                ", 3),  # Spacing
                        (dates, 5)  # Italic
                    ]
                    self.replace_paragraph_multirun_with_metadata(para_idx, text_parts, target_doc)
                    changes.append(f"EDU_INST (para {para_idx})")
                # For Education 2+, skip institution_line (will be handled by degree line below)
            
            if "gpa_paragraph_index" in edu and edu["gpa_paragraph_index"] is not None:
                para_idx = edu["gpa_paragraph_index"]
                gpa_value = edu.get("gpa", "")
                
                # Use multirun: Run 0 is bold label "Cumulative GPA:", Run 1 is normal value
                # Extract label and value
                if ":" in gpa_value:
                    label, value = gpa_value.split(":", 1)
                    text_parts = [
                        (f"{label}:", 0),  # Bold label
                        (value, 1)  # Normal value
                    ]
                    self.replace_paragraph_multirun_with_metadata(para_idx, text_parts, target_doc)
                else:
                    self.replace_paragraph_with_metadata(para_idx, gpa_value, target_doc)
                changes.append(f"EDU_GPA (para {para_idx})")
            
            if "coursework_paragraph_indices" in edu:
                coursework_list = edu.get("coursework", [])
                indices = edu["coursework_paragraph_indices"]
                for idx, coursework_line in enumerate(coursework_list):
                    if idx < len(indices):
                        para_idx = indices[idx]
                        
                        # Use multirun: Run 0 is bold label, Run 1+ is normal value
                        if ":" in coursework_line:
                            label, value = coursework_line.split(":", 1)
                            text_parts = [
                                (f"{label}:", 0),  # Bold label
                                (value, 1)  # Normal value
                            ]
                            self.replace_paragraph_multirun_with_metadata(para_idx, text_parts, target_doc)
                        else:
                            self.replace_paragraph_with_metadata(para_idx, coursework_line, target_doc)
                        changes.append(f"EDU_COURSEWORK (para {para_idx})")
            
            # For Education 2+, write the separate degree line (para 8, etc.)
            if idx > 0 and "institution_line" in edu:
                # Check if there's a separate degree paragraph after the institution
                inst_para = edu["institution_line"]["paragraph_index"]
                degree_para = inst_para + 1  # Degree line is typically next paragraph
                
                # Check if degree paragraph exists and has the right content
                if degree_para < len(target_doc.paragraphs):
                    degree = edu.get("degree", "")
                    dates = edu.get("dates", "")
                    
                    if degree:
                        # Para 8 structure: Run 0 italic degree, Run 1 italic tab, Run 2 normal space, Run 3 italic date
                        text_parts = [
                            (degree, 0),  # Italic
                            ("\t  ", 1),  # Italic tab
                            ("               ", 2),  # Normal spacing
                            (dates, 3)  # Italic
                        ]
                        self.replace_paragraph_multirun_with_metadata(degree_para, text_parts, target_doc)
                        changes.append(f"EDU_DEGREE (para {degree_para})")
        
        # Skills
        for skill_key, skill_data in json_data.get("skills", {}).items():
            para_idx = skill_data["paragraph_index"]
            label = skill_data["label"]
            value = skill_data["value"]
            self.replace_skill_line_with_metadata(para_idx, label, value, target_doc)
            changes.append(f"SKILL_{skill_key.upper()} (para {para_idx})")
        
        # Professional
        for job in json_data.get("professional", []):
            if "header_line" in job:
                header_data = job["header_line"]
                para_idx = header_data["paragraph_index"]
                company = job.get("company", "")
                role = job.get("role", "")
                location = job.get("location", "")
                dates = job.get("dates", "")
                
                # Use multirun to preserve bold/italic formatting
                # Run 0: company (bold), Run 1: role (italic), Run 2: pipe (bold), Run 3: location (normal), Run 4+: spacing+dates (italic)
                text_parts = [
                    (f"{company}, ", 0),  # Bold
                    (f"{role} ", 1),  # Italic
                    ("| ", 2),  # Bold
                    (f"{location}                                                               ", 3),  # Normal
                    (dates, 5)  # Italic
                ]
                self.replace_paragraph_multirun_with_metadata(para_idx, text_parts, target_doc)
                changes.append(f"PROF_HEADER (para {para_idx})")
            
            if "tech_stack" in job and job["tech_stack"]:
                tech_data = job["tech_stack"]
                para_idx = tech_data["paragraph_index"]
                self.replace_skill_line_with_metadata(para_idx, "Tech Stack", tech_data["value"], target_doc)
                changes.append(f"PROF_TECH (para {para_idx})")
            
            for bullet in job.get("bullets", []):
                para_idx = bullet["paragraph_index"]
                self.replace_paragraph_with_metadata(para_idx, bullet["value"], target_doc)
                changes.append(f"PROF_BULLET (para {para_idx})")
        
        # Projects
        for idx, proj in enumerate(json_data.get("projects", [])):
            if "header_line" in proj:
                header_data = proj["header_line"]
                para_idx = header_data["paragraph_index"]
                name = proj.get("name", "")
                role = proj.get("role", "")
                dates = proj.get("dates", "")
                
                # Different projects have different run structures
                # Para 24 (Test-Lab): Run 0 bold name, Run 1 italic role, Run 2 normal spacing, Run 3 italic date
                # Para 28 (3D Printer): Run 0 bold name, Run 1 space, Run 2 italic role, Run 10 italic dates
                # Para 32 (Obstacle): Run 0 bold name, Run 1 italic role+spaces, Run 6 italic dates  
                # Para 36 (Earthworm): Run 0 bold name, Run 1 space, Run 2 italic role, Run 6 italic date
                
                if para_idx == 24:
                    # Test-Lab structure
                    text_parts = [
                        (f"{name}, ", 0),  # Bold
                        (f"{role} ", 1),  # Italic
                        ("                       \t                                                                      ", 2),  # Normal spacing
                        (dates, 3)  # Italic
                    ]
                elif para_idx == 28:
                    # 3D Printer structure
                    text_parts = [
                        (f"{name},", 0),  # Bold
                        (" ", 1),  # Normal space
                        (f"{role}", 2),  # Italic
                        ("\t                       \t\t\t          \t             \t         ", 3),  # Normal+italic spacing
                        (dates, 10)  # Italic dates
                    ]
                elif para_idx == 32:
                    # Obstacle Avoiding Robot structure
                    text_parts = [
                        (f"{name}, ", 0),  # Bold
                        (f"{role}     \t\t\t\t   \t                       ", 1),  # Italic role+tabs
                        (dates, 6)  # Italic dates
                    ]
                elif para_idx == 36:
                    # Earthworm structure
                    text_parts = [
                        (f"{name},", 0),  # Bold
                        (" ", 1),  # Normal space
                        (f"{role}", 2),  # Italic
                        ("\t                       \t                                                ", 3),  # Spacing
                        (dates, 6)  # Italic date
                    ]
                else:
                    # Default structure (fallback)
                    text_parts = [
                        (f"{name}, ", 0),  # Bold
                        (f"{role} ", 1),  # Italic
                        ("                       \t                                                                      ", 2),  # Normal spacing
                        (dates, 3)  # Italic
                    ]
                
                self.replace_paragraph_multirun_with_metadata(para_idx, text_parts, target_doc)
                changes.append(f"PROJ_HEADER (para {para_idx})")
            
            if "tech_stack" in proj and proj["tech_stack"]:
                tech_data = proj["tech_stack"]
                para_idx = tech_data["paragraph_index"]
                # For projects, entire Tech Stack line is italic (including label)
                # Use run 0 (italic) for entire text
                tech_line = f"Tech Stack: {tech_data['value']}"
                self.replace_paragraph_with_metadata(para_idx, tech_line, target_doc)
                changes.append(f"PROJ_TECH (para {para_idx})")
            
            for bullet in proj.get("bullets", []):
                para_idx = bullet["paragraph_index"]
                self.replace_paragraph_with_metadata(para_idx, bullet["value"], target_doc)
                changes.append(f"PROJ_BULLET (para {para_idx})")
        
        # Leadership
        for idx, lead in enumerate(json_data.get("leadership", [])):
            if "header_line" in lead:
                header_data = lead["header_line"]
                para_idx = header_data["paragraph_index"]
                organization = lead.get("organization", "")
                title = lead.get("title", "")
                dates = lead.get("dates", "")
                
                # Leadership entries have different structures, need to check which one
                # First entry: Run 0-1 bold org, Run 2 italic title, Run 4 italic dates
                # Second entry: Run 0 bold org, Run 1 italic title, Run 7 italic dates
                if idx == 0:
                    # First leadership entry structure
                    text_parts = [
                        (f"{organization}, ", 1),  # Bold (run 1)
                        (f"{title}                                                                              ", 2),  # Italic (run 2)
                        (dates, 4)  # Italic (run 4)
                    ]
                else:
                    # Second and other entries structure  
                    text_parts = [
                        (f"{organization}, ", 0),  # Bold (run 0)
                        (f"{title}", 1),  # Italic (run 1)
                        ("        \t\t\t\t                                                          ", 2),  # Normal spacing
                        (dates, 7)  # Italic (run 7)
                    ]
                self.replace_paragraph_multirun_with_metadata(para_idx, text_parts, target_doc)
                changes.append(f"LEAD_HEADER (para {para_idx})")
            
            for bullet in lead.get("bullets", []):
                para_idx = bullet["paragraph_index"]
                self.replace_paragraph_with_metadata(para_idx, bullet["value"], target_doc)
                changes.append(f"LEAD_BULLET (para {para_idx})")
        
        # Save
        target_doc.save(output_path)
        
        print(f"\n✅ Resume generated: {output_path}")
        print(f"📊 Edits applied: {len(changes)}")
        print(f"🎨 Format applied from metadata (margins, spacing, fonts, ALL properties)")
        
        return output_path


# Main execution
if __name__ == '__main__':
    import json
    
    print("🔧 STEP 1: Extract format metadata from original resume")
    original_path = r"C:\Users\dilip\OneDrive\Desktop\ResumeBuilder\reference_docx\resume_optimized_final.docx"
    
    metadata_extractor = FormatMetadata(original_path)
    metadata_extractor.save_metadata("metadata/format_metadata.json")
    
    print("\n🔧 STEP 2: Test building resume with metadata")
    
    # Load content JSON
    with open('templates/resume_content.json', 'r', encoding='utf-8') as f:
        resume_json = json.load(f)
    
    # Test edit
    if "languages" in resume_json["skills"]:
        resume_json["skills"]["languages"]["value"] = "Python, Java, C++, JavaScript, TypeScript, Go, Rust, Kotlin"
    
    # Build with metadata
    builder = EnhancedFormatBuilder(original_path, "metadata/format_metadata.json")
    output = builder.build_resume_from_json(resume_json, "output/test_metadata_output.docx")
    
    print(f"\n✅ Test complete!")
    print(f"📁 Output: {output}")
    print(f"📁 Metadata: metadata/format_metadata.json")
