"""
Cover Letter Format Builder
Uses template-based format preservation like the resume system
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path
import json


class CoverLetterFormatBuilder:
    """Build formatted cover letters from JSON content using template preservation"""
    
    def __init__(self, template_docx, content_json=None):
        """
        Initialize with template DOCX file
        
        Args:
            template_docx: Path to template .docx file
            content_json: Optional content JSON (if None, uses template content)
        """
        self.template_path = template_docx
        self.content = content_json
        
    def build_cover_letter(self, content_json, output_path):
        """
        Build cover letter from JSON content while preserving format
        
        Args:
            content_json: Dict with cover letter content
            output_path: Where to save the generated .docx
            
        Returns:
            str: Path to generated file
        """
        # Load template
        doc = Document(self.template_path)
        
        # Update header (applicant info)
        header_info = content_json.get('header', {})
        recipient_info = content_json.get('recipient', {})
        paragraphs_data = content_json.get('paragraphs', {})
        signature_data = content_json.get('signature', {})
        
        # Clear existing content (keep formatting)
        for paragraph in doc.paragraphs:
            paragraph.clear()
        
        # Rebuild with new content while preserving styles
        para_idx = 0
        
        # 1. Applicant header (right-aligned)
        if para_idx < len(doc.paragraphs):
            p = doc.paragraphs[para_idx]
        else:
            p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run = p.add_run(header_info.get('applicant_name', 'Dilip Kumar'))
        run.bold = True
        run.font.size = Pt(11)
        
        if header_info.get('email'):
            p.add_run('\n' + header_info['email']).font.size = Pt(10)
        if header_info.get('phone'):
            p.add_run('\n' + header_info['phone']).font.size = Pt(10)
        
        para_idx += 1
        
        # 2. Date
        if para_idx < len(doc.paragraphs):
            p = doc.paragraphs[para_idx]
        else:
            p = doc.add_paragraph()
        p.clear()
        
        # Add spacing
        doc.add_paragraph()
        para_idx += 1
        
        # Add date
        if para_idx < len(doc.paragraphs):
            p = doc.paragraphs[para_idx]
        else:
            p = doc.add_paragraph()
        p.clear()
        
        date_text = header_info.get('date', datetime.now().strftime("%B %d, %Y"))
        if date_text == "[Current Date]":
            date_text = datetime.now().strftime("%B %d, %Y")
        p.add_run(date_text).font.size = Pt(11)
        para_idx += 1
        
        # 3. Recipient address
        doc.add_paragraph()  # Space
        para_idx += 1
        
        if para_idx < len(doc.paragraphs):
            p = doc.paragraphs[para_idx]
        else:
            p = doc.add_paragraph()
        p.clear()
        
        hiring_mgr = recipient_info.get('hiring_manager', 'Hiring Manager')
        company = recipient_info.get('company_name', '[Company Name]')
        
        if company and company != '[Company Name]':
            p.add_run(f"{hiring_mgr}\n{company}").font.size = Pt(11)
        else:
            p.add_run(f"To {hiring_mgr}").font.size = Pt(11)
        para_idx += 1
        
        # 4. Salutation
        doc.add_paragraph()  # Space
        para_idx += 1
        
        if para_idx < len(doc.paragraphs):
            p = doc.paragraphs[para_idx]
        else:
            p = doc.add_paragraph()
        p.clear()
        
        salutation = content_json.get('salutation', 'Dear Hiring Manager,')
        p.add_run(salutation).font.size = Pt(11)
        para_idx += 1
        
        # 5. Body paragraphs
        doc.add_paragraph()  # Space before body
        para_idx += 1
        
        paragraph_order = ['opening', 'skills', 'achievements', 'company_knowledge', 'closing']
        
        for para_key in paragraph_order:
            para_content = paragraphs_data.get(para_key, {})
            text = para_content.get('value', '') if isinstance(para_content, dict) else para_content
            
            if text and not text.startswith('['):  # Skip placeholder text
                if para_idx < len(doc.paragraphs):
                    p = doc.paragraphs[para_idx]
                else:
                    p = doc.add_paragraph()
                p.clear()
                
                run = p.add_run(text)
                run.font.size = Pt(11)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para_idx += 1
                
                # Add space between paragraphs (except after last one)
                if para_key != 'closing':
                    doc.add_paragraph()
                    para_idx += 1
        
        # 6. Signature
        doc.add_paragraph()  # Space before signature
        para_idx += 1
        
        if para_idx < len(doc.paragraphs):
            p = doc.paragraphs[para_idx]
        else:
            p = doc.add_paragraph()
        p.clear()
        
        closing_phrase = signature_data.get('closing_phrase', 'Sincerely,')
        p.add_run(closing_phrase + '\n\n').font.size = Pt(11)
        para_idx += 1
        
        if para_idx < len(doc.paragraphs):
            p = doc.paragraphs[para_idx]
        else:
            p = doc.add_paragraph()
        p.clear()
        
        name_run = p.add_run(signature_data.get('name', 'Dilip Kumar'))
        name_run.font.size = Pt(11)
        name_run.bold = True
        
        # Save
        doc.save(output_path)
        return output_path


if __name__ == "__main__":
    # Test with existing template
    template_path = "reference_docx/cover_letter_template.docx"
    
    if not Path(template_path).exists():
        print(f"❌ Template not found: {template_path}")
        print("Please ensure the template exists before running this test.")
        exit(1)
    
    print(f"✅ Using template: {template_path}")
    
    # Test with sample content
    content = {
        "header": {
            "applicant_name": "Dilip Kumar",
            "email": "dilip@example.com",
            "phone": "+1-234-567-8900",
            "date": "[Current Date]"
        },
        "recipient": {
            "hiring_manager": "Hiring Manager",
            "company_name": "TechCorp Inc."
        },
        "salutation": "Dear Hiring Manager,",
        "paragraphs": {
            "opening": {
                "value": "I am writing to express my enthusiasm for the Senior Software Engineer position at TechCorp Inc."
            },
            "skills": {
                "value": "With 5+ years of Python development experience..."
            },
            "achievements": {
                "value": "In my previous role, I led a team that..."
            },
            "company_knowledge": {
                "value": "I am particularly impressed by TechCorp's mission..."
            },
            "closing": {
                "value": "Thank you for considering my application..."
            }
        },
        "signature": {
            "closing_phrase": "Sincerely,",
            "name": "Dilip Kumar"
        }
    }
    builder = CoverLetterFormatBuilder(template_path)
    output = builder.build_cover_letter(content, "test_cover_letter.docx")
    print(f"✅ Test cover letter created: {output}")
