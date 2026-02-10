"""
Simple Cover Letter Builder - Exact Copy of Resume System Pattern
Just copy template and replace text at specific paragraph indices
"""

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path
import json
import shutil


class SimpleCoverLetterBuilder:
    """Simple builder - copy template, replace text, save. Like resume system."""
    
    def __init__(self, metadata_json_path="metadata/cover_letter_format_metadata.json"):
        """Load format metadata"""
        with open(metadata_json_path, 'r', encoding='utf-8') as f:
            self.format_metadata = json.load(f)
    
    def apply_paragraph_format_from_metadata(self, para, para_idx: int):
        """Apply paragraph formatting from metadata"""
        if str(para_idx) not in self.format_metadata['paragraph_formats']:
            return
        
        meta = self.format_metadata['paragraph_formats'][str(para_idx)]
        pf = para.paragraph_format
        
        # Apply alignment
        if meta['alignment']:
            if 'CENTER' in meta['alignment']:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif 'RIGHT' in meta['alignment']:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif 'JUSTIFY' in meta['alignment']:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif 'LEFT' in meta['alignment']:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Apply spacing and indents
        if meta['left_indent']:
            pf.left_indent = meta['left_indent']
        if meta['right_indent']:
            pf.right_indent = meta['right_indent']
        if meta['first_line_indent']:
            pf.first_line_indent = meta['first_line_indent']
        if meta['space_before'] is not None:
            pf.space_before = meta['space_before']
        if meta['space_after'] is not None:
            pf.space_after = meta['space_after']
        if meta['line_spacing']:
            pf.line_spacing = meta['line_spacing']
    
    def apply_run_format_from_metadata(self, run, para_idx: int, run_idx: int = 0):
        """Apply run formatting from metadata"""
        if str(para_idx) not in self.format_metadata['run_formats']:
            return
        
        run_formats = self.format_metadata['run_formats'][str(para_idx)]
        if run_idx >= len(run_formats):
            run_idx = 0
        
        meta = run_formats[run_idx]
        
        # Apply all formatting
        if meta['bold'] is not None:
            run.bold = meta['bold']
        if meta['italic'] is not None:
            run.italic = meta['italic']
        if meta['underline'] is not None:
            run.underline = meta['underline']
        
        if meta['font_name']:
            run.font.name = meta['font_name']
        if meta['font_size']:
            run.font.size = meta['font_size']
        if meta['font_color_rgb']:
            run.font.color.rgb = RGBColor(*meta['font_color_rgb'])
        
        if meta.get('all_caps') is not None:
            run.font.all_caps = meta['all_caps']
        if meta.get('small_caps') is not None:
            run.font.small_caps = meta['small_caps']
    
    def replace_paragraph_text(self, doc: Document, para_idx: int, new_text: str):
        """Replace paragraph text - exact copy of resume system pattern"""
        if para_idx >= len(doc.paragraphs):
            return
        
        para = doc.paragraphs[para_idx]
        
        # Clear existing runs
        for run in para.runs:
            run.text = ""
        
        # Apply paragraph format from metadata
        self.apply_paragraph_format_from_metadata(para, para_idx)
        
        # Add text with formatting
        if para.runs:
            run = para.runs[0]
        else:
            run = para.add_run()
        
        run.text = new_text
        self.apply_run_format_from_metadata(run, para_idx, 0)
    
    def build_from_json(self, content_json: dict, template_path: str, output_path: str):
        """
        Build cover letter from JSON content
        
        Steps (EXACTLY like resume):
        1. Copy template to output
        2. Open copied file
        3. Replace text at specific indices
        4. Save
        """
        # Step 1: Copy template
        shutil.copy(template_path, output_path)
        
        # Step 2: Open document
        doc = Document(output_path)
        
        # Step 3: Replace text at each position
        # Based on paragraph mapping in metadata
        
        # Paragraph 0: Name (header)
        # Paragraph 1: Contact line 1
        # Paragraph 2: Contact line 2
        # These stay as-is from template (have your info)
        
        # Paragraph 3: Date
        date_text = content_json.get('date', datetime.now().strftime("%B %d, %Y"))
        if date_text == "[Date]":
            date_text = datetime.now().strftime("%B %d, %Y")
        self.replace_paragraph_text(doc, 3, date_text)
        
        # Paragraph 4-7: Recipient info (from AI extraction or defaults)
        recipient = content_json.get('recipient', {})
        
        # Hiring manager: use "Hiring Manager" if empty or placeholder
        hiring_mgr = recipient.get('hiring_manager', 'Hiring Manager')
        if not hiring_mgr or hiring_mgr.startswith('['):
            hiring_mgr = 'Hiring Manager'
        self.replace_paragraph_text(doc, 4, hiring_mgr)
        
        # Job title: skip if empty or placeholder
        job_title = recipient.get('job_title', '')
        if job_title and not job_title.startswith('['):
            self.replace_paragraph_text(doc, 5, job_title)
        else:
            self.replace_paragraph_text(doc, 5, '')  # Leave blank
        
        # Company name: use from extraction
        company = recipient.get('company_name', '')
        if company and not company.startswith('['):
            self.replace_paragraph_text(doc, 6, company)
        else:
            self.replace_paragraph_text(doc, 6, '')  # Leave blank
        
        # Company address: leave blank if not available
        address = recipient.get('company_address', '')
        if address and not address.startswith('['):
            self.replace_paragraph_text(doc, 7, address)
        else:
            self.replace_paragraph_text(doc, 7, '')  # Leave blank
        
        # Paragraph 8: Salutation (use Hiring Manager if name not available)
        recipient = content_json.get('recipient', {})
        hiring_mgr = recipient.get('hiring_manager', 'Hiring Manager')
        if not hiring_mgr or hiring_mgr.startswith('['):
            salutation = 'Dear Hiring Manager,'
        else:
            # Use the actual name if available
            salutation = f'Dear {hiring_mgr},'
        self.replace_paragraph_text(doc, 8, salutation)
        
        # Paragraphs 9-13: AI-generated body (5 paragraphs)
        paragraphs = content_json.get('paragraphs', {})
        self.replace_paragraph_text(doc, 9, paragraphs.get('opening', {}).get('value', ''))
        self.replace_paragraph_text(doc, 10, paragraphs.get('skills', {}).get('value', ''))
        self.replace_paragraph_text(doc, 11, paragraphs.get('achievements', {}).get('value', ''))
        self.replace_paragraph_text(doc, 12, paragraphs.get('company_knowledge', {}).get('value', ''))
        self.replace_paragraph_text(doc, 13, paragraphs.get('closing', {}).get('value', ''))
        
        # Paragraph 14: Closing phrase ("Sincerely,")
        # Paragraph 15: Name signature
        # These stay as-is from template
        
        # Step 4: Save
        doc.save(output_path)
        return output_path


# Convenience function for use in generator
def build_cover_letter_simple(ai_paragraphs: dict, company_name: str = None,
                              job_title: str = None, hiring_manager: str = None,
                              output_path: str = None):
    """
    Build cover letter using AI paragraphs
    
    Args:
        ai_paragraphs: Dict with keys: opening, skills, achievements, company_knowledge, closing
        company_name: Extracted company name
        job_title: Extracted job title
        hiring_manager: Extracted hiring manager name
        output_path: Where to save
    """
    # Load content template
    with open("templates/cover_letter_content.json", 'r', encoding='utf-8') as f:
        content = json.load(f)
    
    # Update with AI content
    content['paragraphs']['opening']['value'] = ai_paragraphs.get('opening', '')
    content['paragraphs']['skills']['value'] = ai_paragraphs.get('skills', '')
    content['paragraphs']['achievements']['value'] = ai_paragraphs.get('achievements', '')
    content['paragraphs']['company_knowledge']['value'] = ai_paragraphs.get('company_knowledge', '')
    content['paragraphs']['closing']['value'] = ai_paragraphs.get('closing', '')
    
    # Update recipient info with extracted details
    if company_name:
        content['recipient']['company_name'] = company_name
    if job_title:
        content['recipient']['job_title'] = job_title
    if hiring_manager:
        content['recipient']['hiring_manager'] = hiring_manager
    
    # Update date
    content['date'] = datetime.now().strftime("%B %d, %Y")
    
    # Generate filename if not provided
    if output_path is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"cover_letter_dilip_kumar_{timestamp}.docx"
    
    # Build
    builder = SimpleCoverLetterBuilder()
    builder.build_from_json(
        content,
        "reference_docx/cover_letter_template.docx",
        output_path
    )
    
    return output_path


if __name__ == "__main__":
    # Test
    test_paragraphs = {
        'opening': 'I am excited to apply for the Senior Software Engineer position at Google.',
        'skills': 'With 6+ years of Python development experience, I have built scalable applications.',
        'achievements': 'I led a team that improved system performance by 40% and reduced costs by 30%.',
        'company_knowledge': 'Google\'s commitment to innovation aligns perfectly with my career goals.',
        'closing': 'I look forward to discussing how I can contribute to your team. Thank you for your consideration.'
    }
    
    output = build_cover_letter_simple(test_paragraphs, "Google", "test_simple_cover_letter.docx")
    print(f"✅ Cover letter created: {output}")
