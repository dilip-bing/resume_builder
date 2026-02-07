"""Extract resume content to JSON with proper parsing of headers."""

from docx import Document
import json
import re
from typing import Dict, Any


class ResumeToJSON:
    """Extract resume content to JSON with intelligent parsing."""
    
    def __init__(self, docx_path: str):
        self.doc = Document(docx_path)
        self.docx_path = docx_path
    
    def parse_header_with_dates(self, text: str) -> Dict[str, str]:
        """Parse headers like: 'Name, Role    Dates' or 'Company, Role | Location   Dates'"""
        cleaned = re.sub(r'\s+', ' ', text.strip())
        
        # Extract dates (Month YYYY - Month YYYY) or just (Month YYYY) or (YYYY)
        date_pattern = r'((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}(?:\s*[-–]\s*(?:January|February|March|April|May|June|July|August|September|October|November|December)?\s*\d{4})?)'
        date_match = re.search(date_pattern, cleaned)
        
        if date_match:
            dates = date_match.group(1).strip()
            rest = cleaned[:date_match.start()].strip()
        else:
            # Try just year range
            date_match = re.search(r'(\d{4}\s*[-–]\s*\d{4})', cleaned)
            if date_match:
                dates = date_match.group(1).strip()
                rest = cleaned[:date_match.start()].strip()
            else:
                dates = ""
                rest = cleaned
        
        # Check for location (pipe separator)
        if '|' in rest:
            before_pipe, location = rest.split('|', 1)
            location = location.strip()
        else:
            before_pipe = rest
            location = ""
        
        # Split by comma for first two parts
        if ',' in before_pipe:
            part1, part2 = before_pipe.split(',', 1)
            part1 = part1.strip()
            part2 = part2.strip()
        else:
            part1 = before_pipe.strip()
            part2 = ""
        
        return {
            "part1": part1,
            "part2": part2,
            "location": location,
            "dates": dates
        }
    
    def parse_education_header(self, text: str) -> Dict[str, str]:
        """Parse education header with university, degree, date."""
        # Look for degree patterns
        degree_patterns = [
            r'Master of Science[^E]*',
            r'Bachelor of Engineering[^E]*',
            r'M\.S\.[^E]*',
            r'B\.E\.[^E]*'
        ]
        
        university = text.strip()
        degree = ""
        dates = ""
        
        for pattern in degree_patterns:
            match = re.search(pattern, text)
            if match:
                degree_start = match.start()
                university = text[:degree_start].strip()
                degree_part = text[degree_start:].strip()
                
                # Extract date from degree part
                date_match = re.search(r'(Expected\s+)?([A-Za-z]+\s+\d{4})', degree_part)
                if date_match:
                    dates = date_match.group(0).strip()
                    degree = degree_part[:date_match.start()].strip()
                else:
                    degree = degree_part
                break
        
        return {"university": university, "degree": degree, "dates": dates}
    
    def extract_content_with_positions(self) -> Dict[str, Any]:
        content = {
            "metadata": {
                "source_file": self.docx_path,
                "total_paragraphs": len(self.doc.paragraphs)
            },
            "personal": {},
            "education": [],
            "skills": {},
            "professional": [],
            "projects": [],
            "leadership": []
        }
        
        paras = [(i, p.text.strip()) for i, p in enumerate(self.doc.paragraphs) if p.text.strip()]
        
        # Personal info
        if len(paras) > 0:
            para_idx, text = paras[0]
            content["personal"]["name"] = {"value": text, "paragraph_index": para_idx}
        
        if len(paras) > 1:
            para_idx, text = paras[1]
            parts = [p.strip() for p in text.split('|')]
            content["personal"]["contact_line"] = {
                "value": text,
                "paragraph_index": para_idx,
                "parts": {
                    "location": parts[0] if len(parts) > 0 else "",
                    "phone": parts[1] if len(parts) > 1 else "",
                    "email": parts[2] if len(parts) > 2 else "",
                    "linkedin": parts[3] if len(parts) > 3 else "",
                    "portfolio": parts[4] if len(parts) > 4 else ""
                }
            }
        
        current_section = None
        current_entry = None
        
        for i, (para_idx, text) in enumerate(paras):
            # Section headers
            if text.upper() in ['EDUCATION', 'TECHNICAL SKILLS', 'PROFESSIONAL EXPERIENCE', 'PROJECT EXPERIENCE', 'LEADERSHIP EXPERIENCE']:
                current_section = text.upper()
                current_entry = None
                continue
            
            # EDUCATION
            if current_section == 'EDUCATION':
                if 'University' in text or 'College' in text:
                    parsed = self.parse_education_header(text)
                    edu_entry = {
                        "institution_line": {"value": text, "paragraph_index": para_idx},
                        "university": parsed["university"],
                        "degree": parsed["degree"],
                        "dates": parsed["dates"],
                        "gpa": "",
                        "coursework": [],
                        "gpa_paragraph_index": None,
                        "coursework_paragraph_indices": []
                    }
                    current_entry = edu_entry
                    content["education"].append(current_entry)
                elif current_entry:
                    if 'GPA' in text or 'gpa' in text.lower():
                        current_entry["gpa"] = text
                        current_entry["gpa_paragraph_index"] = para_idx
                    elif 'Coursework' in text or 'coursework' in text.lower():
                        current_entry["coursework"].append(text)
                        current_entry["coursework_paragraph_indices"].append(para_idx)
                    elif ('Bachelor' in text or 'Master' in text or 'Engineering' in text) and not current_entry["degree"]:
                        parsed = self.parse_education_header(text)
                        if parsed["degree"]:
                            current_entry["degree"] = parsed["degree"]
                            current_entry["dates"] = parsed["dates"]
            
            # TECHNICAL SKILLS
            elif current_section == 'TECHNICAL SKILLS':
                if ':' in text:
                    label, value = text.split(':', 1)
                    label = label.strip()
                    value = value.strip()
                    key = label.lower().replace(' ', '_').replace('&', 'and').replace('/', '_')
                    content["skills"][key] = {"value": value, "paragraph_index": para_idx, "label": label}
            
            # PROFESSIONAL EXPERIENCE
            elif current_section == 'PROFESSIONAL EXPERIENCE':
                if 'Tech Stack:' in text:
                    if current_entry:
                        current_entry["tech_stack"] = {"value": text.replace('Tech Stack:', '').strip(), "paragraph_index": para_idx}
                elif text and not text.startswith('●'):
                    # Check if it's a job header
                    has_dates = bool(re.search(r'(January|February|March|April|May|June|July|August|September|October|November|December|20\d{2})', text))
                    if has_dates and (',' in text or '|' in text):
                        parsed = self.parse_header_with_dates(text)
                        job_entry = {
                            "header_line": {"value": text, "paragraph_index": para_idx},
                            "company": parsed["part1"],
                            "role": parsed["part2"],
                            "location": parsed["location"],
                            "dates": parsed["dates"],
                            "tech_stack": None,
                            "bullets": []
                        }
                        current_entry = job_entry
                        content["professional"].append(current_entry)
                    elif current_entry:
                        current_entry["bullets"].append({"value": text, "paragraph_index": para_idx})
                elif current_entry and text.startswith('●'):
                    current_entry["bullets"].append({"value": text, "paragraph_index": para_idx})
            
            # PROJECT EXPERIENCE
            elif current_section == 'PROJECT EXPERIENCE':
                if 'Tech Stack:' in text:
                    if current_entry:
                        current_entry["tech_stack"] = {"value": text.replace('Tech Stack:', '').strip(), "paragraph_index": para_idx}
                elif text and not text.startswith('●'):
                    # Check if it's a project header
                    has_dates = bool(re.search(r'(January|February|March|April|May|June|July|August|September|October|November|December|20\d{2})', text))
                    if has_dates and ',' in text:
                        parsed = self.parse_header_with_dates(text)
                        proj_entry = {
                            "header_line": {"value": text, "paragraph_index": para_idx},
                            "name": parsed["part1"],
                            "role": parsed["part2"],
                            "dates": parsed["dates"],
                            "tech_stack": None,
                            "bullets": []
                        }
                        current_entry = proj_entry
                        content["projects"].append(current_entry)
                    elif current_entry:
                        current_entry["bullets"].append({"value": text, "paragraph_index": para_idx})
                elif current_entry and text.startswith('●'):
                    current_entry["bullets"].append({"value": text, "paragraph_index": para_idx})
            
            # LEADERSHIP EXPERIENCE
            elif current_section == 'LEADERSHIP EXPERIENCE':
                if text and not text.startswith('●'):
                    # Check if it's a leadership header
                    has_dates = bool(re.search(r'(January|February|March|April|May|June|July|August|September|October|November|December|20\d{2})', text))
                    if has_dates and ',' in text:
                        parsed = self.parse_header_with_dates(text)
                        lead_entry = {
                            "header_line": {"value": text, "paragraph_index": para_idx},
                            "organization": parsed["part1"],
                            "title": parsed["part2"],
                            "dates": parsed["dates"],
                            "bullets": []
                        }
                        current_entry = lead_entry
                        content["leadership"].append(current_entry)
                    elif current_entry:
                        current_entry["bullets"].append({"value": text, "paragraph_index": para_idx})
                elif current_entry and text.startswith('●'):
                    current_entry["bullets"].append({"value": text, "paragraph_index": para_idx})
        
        return content
    
    def save_to_json(self, output_path: str = "resume_content.json"):
        content = self.extract_content_with_positions()
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(content, f, indent=2, ensure_ascii=False)
        print(f"✅ Saved to: {output_path}")
        self.print_summary(content)
        return output_path
    
    def print_summary(self, content: Dict[str, Any]):
        print(f"\n📊 Extraction Summary:")
        print(f"   Personal fields: {len(content['personal'])}")
        print(f"   Education entries: {len(content['education'])}")
        print(f"   Skill categories: {len(content['skills'])}")
        print(f"   Professional jobs: {len(content['professional'])}")
        print(f"   Projects: {len(content['projects'])}")
        print(f"   Leadership roles: {len(content['leadership'])}")


if __name__ == '__main__':
    resume_path = r"C:\Users\dilip\OneDrive\Desktop\ResumeBuilder\reference_docx\resume_optimized_final.docx"
    extractor = ResumeToJSON(resume_path)
    extractor.save_to_json("templates/resume_content.json")
