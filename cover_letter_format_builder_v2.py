"""
Cover Letter Format Builder with Metadata-based Format Preservation
Works exactly like the resume system - uses format_metadata.json for exact formatting
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path
import json
import shutil


class CoverLetterFormatBuilder:
    """Build formatted cover letters using metadata-based format preservation"""
    
    def __init__(self, template_docx_path: str, metadata_json_path: str = None):
        """
        Initialize with template and metadata
        
        Args:
            template_docx_path: Path to cover letter template DOCX
            metadata_json_path: Path to format metadata JSON (auto-inferred if None)
        """
        self.template_path = template_docx_path
        
        if metadata_json_path is None:
            metadata_json_path = "metadata/cover_letter_format_metadata.json"
        
        self.metadata_path = metadata_json_path
        
        # Load format metadata
        with open(self.metadata_path, 'r', encoding='utf-8') as f:
            self.format_metadata = json.load(f)
    
    def apply_paragraph_format_from_metadata(self, para, para_idx: int):
        """Apply paragraph formatting from metadata"""
        if str(para_idx) not in self.format_metadata['paragraph_formats']:
            return
        
        meta = self.format_metadata['paragraph_formats'][str(para_idx)]
        pf = para.paragraph_format
        
        # Apply alignment
        if meta['alignment']:
            if 'CENTER' in meta['alignment']:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif 'RIGHT' in meta['alignment']:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif 'LEFT' in meta['alignment']:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif 'JUSTIFY' in meta['alignment']:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Apply indents and spacing
        if meta['left_indent']:
            pf.left_indent = meta['left_indent']
        if meta['right_indent']:
            pf.right_indent = meta['right_indent']
        if meta['first_line_indent']:
            pf.first_line_indent = meta['first_line_indent']
        if meta['space_before']:
            pf.space_before = meta['space_before']
        if meta['space_after'] is not None:
            pf.space_after = meta['space_after']
        if meta['line_spacing']:
            pf.line_spacing = meta['line_spacing']
    
    def apply_run_format_from_metadata(self, run, para_idx: int, run_idx: int = 0):
        """Apply run formatting from metadata"""
        if str(para_idx) not in self.format_metadata['run_formats']:
            return
        
        run_formats = self.format_metadata['run_formats'][str(para_idx)]
        if run_idx >= len(run_formats):
            run_idx = 0  # Use first run format as fallback
        
        meta = run_formats[run_idx]
        
        # Apply formatting
        if meta['bold'] is not None:
            run.bold = meta['bold']
        if meta['italic'] is not None:
            run.italic = meta['italic']
        if meta['underline'] is not None:
            run.underline = meta['underline']
        
        if meta['font_name']:
            run.font.name = meta['font_name']
        if meta['font_size']:
            run.font.size = meta['font_size']
        if meta['font_color_rgb']:
            run.font.color.rgb = RGBColor(*meta['font_color_rgb'])
        
        if meta.get('all_caps') is not None:
            run.font.all_caps = meta['all_caps']
        if meta.get('small_caps') is not None:
            run.font.small_caps = meta['small_caps']
    
    def replace_paragraph_text(self, doc: Document, para_idx: int, new_text: str):
        """Replace paragraph text while preserving format from metadata"""
        if para_idx >= len(doc.paragraphs):
            return
        
        para = doc.paragraphs[para_idx]
        
        # Clear existing runs
        for run in para.runs:
            run.text = ""
        
        # Apply paragraph format
        self.apply_paragraph_format_from_metadata(para, para_idx)
        
        # Add new text with formatting
        if para.runs:
            run = para.runs[0]
        else:
            run = para.add_run()
        
        run.text = new_text
        
        # Apply run format
        self.apply_run_format_from_metadata(run, para_idx, 0)
    
    def build_cover_letter_from_content(self, content_json: dict, output_path: str):
        """
        Build cover letter from content JSON using template and metadata
        
        Args:
            content_json: Dict with cover letter content
            output_path: Where to save generated DOCX
        
        Returns:
            str: Path to generated file
        """
        # Copy template to output
        shutil.copy(self.template_path, output_path)
        
        # Load the copied document
        doc = Document(output_path)
        
        # Get paragraph mapping from content
        mapping = content_json.get('_metadata', {}).get('paragraph_mapping', {})
        
        # Replace each paragraph
        for para_idx_str, json_path in mapping.items():
            para_idx = int(para_idx_str)
            
            # Get value from content JSON
            value = self._get_nested_value(content_json, json_path)
            
            if value and not value.startswith('['):  # Skip placeholders
                # Special handling for date
                if json_path == 'date':
                    if value == '[Date]':
                        value = datetime.now().strftime("%B %d, %Y")
                
                # Replace paragraph text
                self.replace_paragraph_text(doc, para_idx, value)
        
        # Save
        doc.save(output_path)
        return output_path
    
    def _get_nested_value(self, data: dict, path: str):
        """Get value from nested dict using dot notation"""
        keys = path.split('.')
        value = data
        
        for key in keys:
            if isinstance(value, dict):
                value = value.get(key, '')
            else:
                return ''
        
        return value if value else ''


def build_cover_letter(job_description: str, resume_text: str = "", 
                       context: str = "", output_path: str = None):
    """
    Convenience function to build cover letter with AI-generated content
    
    Args:
        job_description: Job posting text
        resume_text: Optimized resume content
        context: Additional context
        output_path: Where to save (auto-generated if None)
    
    Returns:
        tuple: (output_path, company_name)
    """
    from cover_letter_generator import CoverLetterGenerator
    import os
    
    # Load content template
    template_json_path = "templates/cover_letter_content.json"
    with open(template_json_path, 'r', encoding='utf-8') as f:
        content = json.load(f)
    
    # Initialize AI generator
    api_key = os.getenv("GEMINI_API_KEY")
    generator = CoverLetterGenerator(api_key=api_key)
    
    # Generate paragraphs
    result = generator.generate_cover_letter(job_description, resume_text, context)
    company_name = result['company_name']
    paragraphs = result['paragraphs']
    
    # Update content with AI-generated text
    content['paragraphs']['opening']['value'] = paragraphs.get('opening', '')
    content['paragraphs']['skills']['value'] = paragraphs.get('skills', '')
    content['paragraphs']['achievements']['value'] = paragraphs.get('achievements', '')
    content['paragraphs']['company_knowledge']['value'] = paragraphs.get('company_knowledge', '')
    content['paragraphs']['closing']['value'] = paragraphs.get('closing', '')
    
    # Update recipient info
    if company_name:
        content['recipient']['company_name'] = company_name
    
    content['date'] = datetime.now().strftime("%B %d, %Y")
    
    # Generate output filename
    if output_path is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"cover_letter_dilip_kumar_{timestamp}.docx"
    
    # Build cover letter
    builder = CoverLetterFormatBuilder(
        "reference_docx/cover_letter_template.docx",
        "metadata/cover_letter_format_metadata.json"
    )
    
    builder.build_cover_letter_from_content(content, output_path)
    
    return output_path, company_name


if __name__ == "__main__":
    # Test
    test_content = {
        "header": {
            "name": "Dilip Kumar Thirukonda Chandrasekaran",
            "contact_line_1": "(607) 624-9390 | dthirukondac@binghamton.edu",
            "contact_line_2": "linkedin.com/dilipkumartc | dilip-bing.github.io/portfolio/"
        },
        "date": datetime.now().strftime("%B %d, %Y"),
        "recipient": {
            "hiring_manager": "Hiring Manager",
            "job_title": "Senior Software Engineer",
            "company_name": "TechCorp Inc.",
            "company_address": "123 Tech Street, San Francisco, CA 94105"
        },
        "salutation": "Dear Hiring Manager,",
        "paragraphs": {
            "opening": {
                "value": "I am writing to express my strong interest in the Senior Software Engineer position at TechCorp Inc. With over six years of experience in Python development and a proven track record of delivering scalable solutions, I am excited about the opportunity to contribute to your innovative team."
            },
            "skills": {
                "value": "Throughout my career, I have developed expertise in building robust web applications using Django and Flask frameworks. My experience includes designing RESTful APIs, implementing microservices architectures, and optimizing database performance. At my current role, I successfully led the development of a high-traffic application serving over 1 million users, improving response times by 40%."
            },
            "achievements": {
                "value": "One of my key accomplishments was architecting and implementing a real-time data processing pipeline that reduced data latency from hours to minutes, enabling faster business decisions. I also spearheaded the migration of legacy systems to AWS cloud infrastructure, resulting in 30% cost savings and improved system reliability. These experiences have equipped me with the skills to tackle complex technical challenges."
            },
            "company_knowledge": {
                "value": "I am particularly drawn to TechCorp's commitment to innovation and your recent advancements in AI-driven solutions. Your company's focus on creating technology that makes a meaningful impact aligns perfectly with my professional goals. I am impressed by your collaborative culture and emphasis on continuous learning, values that I share and have practiced throughout my career."
            },
            "closing": {
                "value": "I am enthusiastic about the possibility of bringing my technical expertise and passion for innovation to TechCorp Inc. I would welcome the opportunity to discuss how my background and skills can contribute to your team's success. Thank you for considering my application, and I look forward to the possibility of speaking with you soon."
            }
        },
        "signature": {
            "closing_phrase": "Sincerely,",
            "name": "Dilip Kumar Thirukonda Chandrasekaran"
        },
        "_metadata": {
            "template_file": "reference_docx/cover_letter_template.docx",
            "format_metadata_file": "metadata/cover_letter_format_metadata.json",
            "paragraph_mapping": {
                "0": "header.name",
                "1": "header.contact_line_1",
                "2": "header.contact_line_2",
                "3": "date",
                "4": "recipient.hiring_manager",
                "5": "recipient.job_title",
                "6": "recipient.company_name",
                "7": "recipient.company_address",
                "8": "salutation",
                "9": "paragraphs.opening.value",
                "10": "paragraphs.skills.value",
                "11": "paragraphs.achievements.value",
                "12": "paragraphs.company_knowledge.value",
                "13": "paragraphs.closing.value",
                "14": "signature.closing_phrase",
                "15": "signature.name"
            }
        }
    }
    
    builder = CoverLetterFormatBuilder(
        "reference_docx/cover_letter_template.docx",
        "metadata/cover_letter_format_metadata.json"
    )
    
    output = builder.build_cover_letter_from_content(test_content, "test_cover_letter_metadata.docx")
    print(f"✅ Test cover letter created: {output}")
